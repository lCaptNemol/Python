#Strip Doc Plus Convert to PDF

from docx import Document
import os
from pdf2docx import Converter

def remove_text_tables_and_images(input_directory, output_directory, pdf_output_directory):
    text_to_remove = "This document was exported from DFE. Any edits made during review must be copied back into DFE and follow its content structures and best practices."

    # Ensure output directories exist
    os.makedirs(output_directory, exist_ok=True)
    os.makedirs(pdf_output_directory, exist_ok=True)

    for filename in os.listdir(input_directory):
        if filename.startswith("~$") or not filename.endswith(".docx"):
            continue  # Skip temporary or non-Word files

        doc_path = os.path.join(input_directory, filename)
        output_path = os.path.join(output_directory, filename)
        pdf_output_path = os.path.join(pdf_output_directory, filename.replace(".docx", ".pdf"))

        try:
            doc = Document(doc_path)

            # ✅ **Step 1: Remove tables containing the specific text**
            tables_to_remove = []
            for i, table in enumerate(doc.tables):
                if any(text_to_remove in cell.text for row in table.rows for cell in row.cells):
                    tables_to_remove.append(i)
            for i in sorted(tables_to_remove, reverse=True):
                tbl = doc.tables[i]._element
                tbl.getparent().remove(tbl)

            # ✅ **Step 2: Remove tables with "Audience" in the first row**
            tables_to_remove = []
            for i, table in enumerate(doc.tables):
                if table.rows and any(cell.text.strip().lower() == "audience" for cell in table.rows[0].cells):
                    tables_to_remove.append(i)
            for i in sorted(tables_to_remove, reverse=True):
                tbl = doc.tables[i]._element
                tbl.getparent().remove(tbl)

            # ✅ **Step 3: Remove all images from the document**
            for para in doc.paragraphs:
                for run in para.runs:
                    if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing") or \
                       run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"):
                        run._element.getparent().remove(run._element)

            # ✅ **Step 4: Remove images inside tables**
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing") or \
                                   run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"):
                                    run._element.getparent().remove(run._element)

            # Save the cleaned Word document
            doc.save(output_path)
            print(f"✅ Processed: {filename}")

            # ✅ **Step 5: Convert cleaned Word document to PDF**
            cv = Converter(output_path)
            cv.convert(pdf_output_path, start=0, end=None)  # Convert entire document
            cv.close()
            print(f"📄 Converted to PDF: {filename.replace('.docx', '.pdf')}")

        except Exception as e:
            print(f"❌ Error processing {filename}: {e}")

# Example usage
input_directory = "/Users/km/Documents/Projects/Combine_Word_Docs/Convert"
output_directory = "/Users/km/Documents/Projects/Combine_Word_Docs/Clean"
pdf_output_directory = "/Users/km/Documents/Projects/Combine_Word_Docs/PDFs"

remove_text_tables_and_images(input_directory, output_directory, pdf_output_directory)

