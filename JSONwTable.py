import docx
import json
import os
import re

def extract_table(table):
    """Converts a Word table into a list of dictionaries (rows)."""
    rows = table.rows
    if not rows:
        return []

    headers = [cell.text.strip() for cell in rows[0].cells]
    table_data = []

    for row in rows[1:]:  # Skip header row
        row_data = {headers[i]: cell.text.strip() for i, cell in enumerate(row.cells)}
        table_data.append(row_data)

    return table_data

def is_list_item(text):
    """Checks if a paragraph is part of a list (manual or Word-style)."""
    return bool(re.match(r"^(-|\*|•|\d+\.|[a-z]\.)\s+", text))  # Detects "- ", "* ", "• ", "1. ", "a. "

def clean_list_item(text):
    """Removes list markers like '- ', '• ', '1. ', etc."""
    return re.sub(r"^(-|\*|•|\d+\.|[a-z]\.)\s+", "", text).strip()

def docx_to_json(docx_path, output_path):
    """Converts a Word document (headings, paragraphs, lists, and tables) to JSON."""
    document_data = []
    current_section = None
    current_subsection = None
    current_subsubsection = None
    current_list = None  # Track list items

    try:
        doc = docx.Document(docx_path)

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            style_name = paragraph.style.name if paragraph.style else ""

            if not text:
                continue

            # Handle headings
            if style_name.startswith("Heading 1"):
                current_section = {"title": text, "content": []}
                document_data.append(current_section)
                current_subsection = None
                current_subsubsection = None
                current_list = None  # Reset list tracking
            elif style_name.startswith("Heading 2"):
                current_subsection = {"title": text, "content": []}
                if current_section:
                    current_section["content"].append(current_subsection)
                current_subsubsection = None
                current_list = None
            elif style_name.startswith("Heading 3"):
                current_subsubsection = {"title": text, "content": []}
                if current_subsection:
                    current_subsection["content"].append(current_subsubsection)
                current_list = None
            # Handle lists
            elif is_list_item(text):
                if current_list is None:
                    current_list = {"type": "list", "items": []}
                    if current_subsubsection:
                        current_subsubsection["content"].append(current_list)
                    elif current_subsection:
                        current_subsection["content"].append(current_list)
                    elif current_section:
                        current_section["content"].append(current_list)

                list_item_text = clean_list_item(text)  # Remove bullet points
                current_list["items"].append({"text": list_item_text})
            else:  # Normal paragraph
                if current_list:
                    current_list = None  # End list tracking

                content_item = {"type": "paragraph", "text": text}

                if current_subsubsection:
                    current_subsubsection["content"].append(content_item)
                elif current_subsection:
                    current_subsection["content"].append(content_item)
                elif current_section:
                    current_section["content"].append(content_item)
                else:
                    if not document_data:
                        current_section = {"title": "Introduction", "content": []}
                        document_data.append(current_section)
                    current_section["content"].append(content_item)

        # Process tables
        for table in doc.tables:
            table_data = extract_table(table)

            if table_data:
                table_item = {"type": "table", "data": table_data}

                if current_subsubsection:
                    current_subsubsection["content"].append(table_item)
                elif current_subsection:
                    current_subsection["content"].append(table_item)
                elif current_section:
                    current_section["content"].append(table_item)
                else:
                    if not document_data:
                        current_section = {"title": "Introduction", "content": []}
                        document_data.append(current_section)
                    current_section["content"].append(table_item)

        # Convert to JSON and save
        json_output = json.dumps(document_data, indent=4, ensure_ascii=False)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(json_output)

        print(f"Conversion successful. JSON output saved to {output_path}")

    except docx.opc.exceptions.PackageNotFoundError:
        print(f"Error: File not found: {docx_path}")
    except Exception as e:
        print(f"An error occurred: {e} while processing {docx_path}")

def process_directory(input_dir, output_dir):
    """Processes all .docx files in the input directory and converts them to JSON."""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for filename in os.listdir(input_dir):
        if filename.endswith(".docx") and not filename.startswith("~$"):  # Ignore temporary Word files
            input_path = os.path.join(input_dir, filename)
            output_filename = os.path.splitext(filename)[0] + ".json"
            output_path = os.path.join(output_dir, output_filename)

            docx_to_json(input_path, output_path)

# ==== SET YOUR INPUT AND OUTPUT DIRECTORY HERE ====
input_directory = "/Users/km/Documents/Convert"  # Change this to your input folder
output_directory = "/Users/km/Documents/JSON"    # Change this to your output folder

process_directory(input_directory, output_directory)