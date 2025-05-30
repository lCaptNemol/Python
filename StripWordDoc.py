from docx import Document
import os
from lxml import etree

def remove_text_tables_and_images(input_directory, output_directory):
    text_to_remove = "This document was exported from DFE. Any edits made during review must be copied back into DFE and follow its content structures and best practices."

    # Ensure output directory exists
    os.makedirs(output_directory, exist_ok=True)

    for filename in os.listdir(input_directory):
        # Skip temporary Word files (~$)
        if filename.startswith("~$") or not filename.endswith(".docx"):
            continue

        doc_path = os.path.join(input_directory, filename)
        output_path = os.path.join(output_directory, filename)

        try:
            doc = Document(doc_path)

            # ✅ **Step 1: Identify tables containing the exact text and remove them**
            tables_to_remove = []
            for i, table in enumerate(doc.tables):
                if any(text_to_remove in cell.text for row in table.rows for cell in row.cells):
                    tables_to_remove.append(i)

            for i in sorted(tables_to_remove, reverse=True):
                tbl = doc.tables[i]._element
                tbl.getparent().remove(tbl)  # Remove from the document XML structure

            # ✅ **Step 2: Identify and remove tables with "Audience" in the first row**
            tables_to_remove = []
            for i, table in enumerate(doc.tables):
                if table.rows and any(cell.text.strip().lower() == "audience" for cell in table.rows[0].cells):
                    tables_to_remove.append(i)

            for i in sorted(tables_to_remove, reverse=True):
                tbl = doc.tables[i]._element
                tbl.getparent().remove(tbl)

            # ✅ **Step 3: Remove all images from the document**
            for para in doc.paragraphs:
                for run in para.runs:
                    if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing") or \
                       run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"):
                        run._element.getparent().remove(run._element)  # Remove image

            # ✅ **Step 4: Remove images inside tables**
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing") or \
                                   run._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"):
                                    run._element.getparent().remove(run._element)  # Remove image

            # Save the modified document
            doc.save(output_path)
            print(f"✅ Processed: {filename}")

        except Exception as e:
            print(f"❌ Error processing {filename}: {e}")

# Example usage
input_directory = "/Users/km/Documents/Projects/Combine_Word_Docs/Convert"
output_directory = "/Users/km/Documents/Projects/Combine_Word_Docs/Clean"
remove_text_tables_and_images(input_directory, output_directory)